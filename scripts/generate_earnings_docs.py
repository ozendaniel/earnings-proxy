#!/usr/bin/env python
"""Generate one Word (.docx) per company-quarter using the earnings-proxy Render endpoint.

Workflow:
- Read targets from a CSV (symbol,quarter)
- Call /summary?symbol=...&quarter=... with x-action-key
- Write a .docx file into a local Dropbox folder (syncs automatically)

Usage examples:
  py generate_earnings_docs.py --targets targets.csv --action-key YOUR_ACTION_API_KEY
  py generate_earnings_docs.py --targets targets.csv --action-key-env ACTION_API_KEY

Optional:
  --base-url https://earnings-proxy.onrender.com
  --dropbox-dir "C:\\Users\\Dan\\Dropbox"
  --out-subdir "Earnings Summaries"

Notes:
- This script writes the API response 'markdown' as plain text into the doc.
- No secrets are stored on disk unless you choose to.
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import re
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

import requests
from docx import Document


DEFAULT_BASE_URL = "https://earnings-proxy.onrender.com"


@dataclass(frozen=True)
class Target:
    symbol: str
    quarter: str


def die(msg: str, code: int = 1) -> None:
    print(f"ERROR: {msg}", file=sys.stderr)
    raise SystemExit(code)


def normalize_symbol(s: str) -> str:
    return s.strip().upper()


_QUARTER_RE = re.compile(r"^\d{4}Q[1-4]$")


def normalize_quarter(q: str) -> str:
    q = q.strip().upper()
    if not _QUARTER_RE.match(q):
        die(f"Quarter must look like 2024Q4 (got '{q}')")
    return q


def safe_filename(name: str) -> str:
    # Windows-illegal: < > : " / \ | ? * plus control chars
    # Replace characters that are invalid in Windows filenames
    name = re.sub(r"[<>:\"/\\|?*]", "-", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name


def read_targets_csv(path: Path) -> List[Target]:
    if not path.exists():
        die(f"Targets file not found: {path}")

    out: List[Target] = []
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        if not reader.fieldnames:
            die("targets.csv is missing a header row. Expected: symbol,quarter")
        # Accept a few common variants
        fields = {h.strip().lower(): h for h in reader.fieldnames}
        if "symbol" not in fields or "quarter" not in fields:
            die(f"targets.csv must have columns symbol,quarter (got {reader.fieldnames})")

        for row in reader:
            sym = normalize_symbol(row[fields["symbol"]] or "")
            qtr = normalize_quarter(row[fields["quarter"]] or "")
            if sym:
                out.append(Target(sym, qtr))

    if not out:
        die("No targets found in CSV.")
    return out


def resolve_dropbox_dir(explicit: Optional[str]) -> Path:
    if explicit:
        p = Path(explicit).expanduser()
        if not p.exists():
            die(f"Dropbox dir does not exist: {p}")
        return p

    env = os.environ.get("DROPBOX_DIR")
    if env:
        p = Path(env).expanduser()
        if p.exists():
            return p

    # Common default on Windows
    userprofile = os.environ.get("USERPROFILE")
    if userprofile:
        p = Path(userprofile) / "Dropbox"
        if p.exists():
            return p

    die(
        "Could not find your Dropbox folder. Provide --dropbox-dir or set DROPBOX_DIR env var."
    )
    raise AssertionError("unreachable")


def fetch_summary(
    *,
    base_url: str,
    action_key: str,
    symbol: str,
    quarter: str,
    timeout_s: int = 60,
    retries: int = 3,
) -> dict:
    url = f"{base_url.rstrip('/')}/summary"
    params = {"symbol": symbol, "quarter": quarter}
    headers = {"x-action-key": action_key}

    last_err: Optional[Exception] = None
    for attempt in range(retries + 1):
        try:
            r = requests.get(url, params=params, headers=headers, timeout=timeout_s)
            if r.status_code == 200:
                return r.json()

            # Retry on common transient statuses
            if r.status_code in (429, 502, 503, 504):
                raise RuntimeError(f"HTTP {r.status_code}: {r.text}")

            # Non-retryable
            raise RuntimeError(f"HTTP {r.status_code}: {r.text}")

        except Exception as e:
            last_err = e
            if attempt < retries:
                sleep_s = 2 ** attempt
                print(f"Retrying {symbol} {quarter} in {sleep_s}s بسبب: {e}")
                time.sleep(sleep_s)
            else:
                break

    die(f"Failed to fetch summary for {symbol} {quarter}: {last_err}")
    raise AssertionError("unreachable")


def write_docx(*, out_path: Path, title: str, body_text: str, meta: dict) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=1)

    # Add a short metadata block
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')} ")
    if meta.get("source"):
        doc.add_paragraph(f"Source: {meta['source']}")

    doc.add_paragraph("")
    for line in body_text.splitlines():
        # Keep markdown as plain text; better formatting can be added later.
        doc.add_paragraph(line)

    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate earnings summary .docx files into Dropbox.")
    ap.add_argument("--targets", default="targets.csv", help="Path to CSV with symbol,quarter")
    ap.add_argument("--base-url", default=DEFAULT_BASE_URL, help="Base URL for earnings-proxy")

    key_group = ap.add_mutually_exclusive_group(required=True)
    key_group.add_argument("--action-key", help="ACTION_API_KEY value")
    key_group.add_argument(
        "--action-key-env",
        help="Name of env var that holds ACTION_API_KEY (recommended). Example: ACTION_API_KEY",
    )

    ap.add_argument(
        "--dropbox-dir",
        default=None,
        help='Dropbox folder path (e.g., "C:\\Users\\ozend\\Dropbox"). If omitted, tries DROPBOX_DIR then USERPROFILE\\Dropbox.',
    )
    ap.add_argument(
        "--out-subdir",
        default="O3 Industries\\#Automated Transcript Summaries",
        help="Subfolder under Dropbox",
    )

    args = ap.parse_args()

    action_key = args.action_key
    if args.action_key_env:
        action_key = os.environ.get(args.action_key_env)
        if not action_key:
            die(f"Env var {args.action_key_env} is not set")

    targets_path = Path(args.targets)
    targets = read_targets_csv(targets_path)

    dropbox_dir = resolve_dropbox_dir(args.dropbox_dir)
    out_dir = dropbox_dir / args.out_subdir

    print(f"Using base URL: {args.base_url}")
    print(f"Reading targets: {targets_path.resolve()}")
    print(f"Writing docs to: {out_dir.resolve()}")
    print("")

    for t in targets:
        print(f"Fetching: {t.symbol} {t.quarter} ...")
        payload = fetch_summary(
            base_url=args.base_url,
            action_key=action_key,
            symbol=t.symbol,
            quarter=t.quarter,
        )

        md = payload.get("markdown") or ""
        if not md:
            die(f"No 'markdown' field returned for {t.symbol} {t.quarter}: {json.dumps(payload)[:500]}")

        filename = safe_filename(f"{t.symbol}_{t.quarter}.docx")
        out_path = out_dir / filename

        title = f"{t.symbol} — {t.quarter}"
        write_docx(out_path=out_path, title=title, body_text=md, meta={"source": args.base_url})
        print(f"Wrote: {out_path}")

    print("\nDone.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
